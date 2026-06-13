from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
TABLE_FILL = "F2F4F7"
BORDER = "DADCE0"


def set_run_font(run, *, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    font = run.font
    font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_border_bottom(paragraph, color: str = "B7C4D6", size: str = "8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def patch_list_style(doc: Document, style_name: str):
    style = doc.styles[style_name]
    fmt = style.paragraph_format
    fmt.left_indent = Inches(0.5)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_after = Pt(4)
    fmt.line_spacing = 1.25


def configure_document(doc: Document, title: str, subtitle: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    patch_list_style(doc, "List Bullet")
    patch_list_style(doc, "List Number")

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    left = header_p.add_run("Hirewave")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header_p.add_run(" | ")
    right = header_p.add_run(subtitle)
    set_run_font(right, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_p.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer_p._p.append(field_begin)
    footer_p._p.append(instr)
    footer_p._p.append(field_end)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=22, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=11, color=MUTED)
    paragraph_border_bottom(p)


def inline_runs(paragraph, text: str, *, bold_default=False):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=10, color=DARK_BLUE, bold=False)
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, bold=bold_default)


def add_markdown_table(doc: Document, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [cell.strip().replace("\\`", "`") for cell in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) < 2:
        return
    data_rows = [parsed[0], *parsed[2:]]
    col_count = max(len(row) for row in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = "Table Grid"
    widths = column_widths(data_rows, col_count)
    set_table_width(table, widths)
    set_cell_margins(table)

    for row_idx, cells in enumerate(data_rows):
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            text = cells[col_idx] if col_idx < len(cells) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            inline_runs(paragraph, text, bold_default=row_idx == 0)
            for run in paragraph.runs:
                if row_idx == 0:
                    run.bold = True
            if row_idx == 0:
                set_cell_fill(cell, TABLE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def column_widths(rows: list[list[str]], col_count: int) -> list[int]:
    available = 9360
    if col_count <= 1:
        return [available]
    lengths = []
    for idx in range(col_count):
        max_len = max((len(row[idx]) if idx < len(row) else 0) for row in rows)
        lengths.append(max(10, min(max_len, 70)))
    total = sum(lengths)
    widths = [max(1100, int(available * length / total)) for length in lengths]
    delta = available - sum(widths)
    widths[-1] += delta
    return widths


def add_code_block(doc: Document, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, size=9, color=DARK_BLUE)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")


def add_markdown(doc: Document, markdown: str, *, source_name: str | None = None):
    if source_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"Source: {source_name}")
        set_run_font(run, size=9, color=MUTED, bold=True)

    lines = markdown.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    while idx < len(lines):
        line = lines[idx].rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_table()
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            idx += 1
            continue
        flush_table()

        if not line.strip():
            idx += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            text = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(text, style="Heading 1")
            elif level == 2:
                p = doc.add_paragraph(text, style="Heading 2")
            else:
                p = doc.add_paragraph(text, style="Heading 3")
            p.paragraph_format.keep_with_next = True
            idx += 1
            continue

        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            inline_runs(p, bullet.group(1).strip())
            idx += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            inline_runs(p, numbered.group(1).strip())
            idx += 1
            continue

        p = doc.add_paragraph()
        inline_runs(p, line)
        idx += 1

    flush_table()
    if code_lines:
        add_code_block(doc, code_lines)


def build_docx(title: str, subtitle: str, sources: list[Path], output: Path):
    doc = Document()
    configure_document(doc, title, subtitle)
    for source_idx, source in enumerate(sources):
        if source_idx:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_markdown(doc, source.read_text(encoding="utf-8"), source_name=str(source.relative_to(ROOT)))
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    build_docx(
        "Hirewave API Reference",
        "MVP route contracts, telemetry side effects, and sandbox metadata",
        [ROOT / "docs" / "07-api-reference.md"],
        OUT_DIR / "Hirewave_API_Reference.docx",
    )
    build_docx(
        "Hirewave Product And Technical Specification",
        "Product requirements, architecture, route inventory, and AI assessment strategy",
        [
            ROOT / "docs" / "02-product-requirements.md",
            ROOT / "docs" / "03-technical-architecture.md",
            ROOT / "docs" / "05-routes-and-features.md",
            ROOT / "docs" / "06-enterprise-ai-assessment-strategy.md",
        ],
        OUT_DIR / "Hirewave_Product_Technical_Specification.docx",
    )


if __name__ == "__main__":
    main()
